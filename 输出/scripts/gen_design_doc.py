# -*- coding: utf-8 -*-
"""生成系统设计说明书 docx"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

DOC = Document()

style = DOC.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for i, (sz, color) in enumerate([(20, '1F4E79'), (15, '1F4E79'), (12.5, '2E74B5'), (11, '2E74B5')], start=1):
    h = DOC.styles[f'Heading {i}']
    h.font.name = 'Calibri'
    h.font.size = Pt(sz)
    h.font.color.rgb = RGBColor.from_string(color)
    h.font.bold = True
    h.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def set_east_asia(run):
    run.font.name = 'Calibri'
    rpr = run._element.get_or_add_rPr()
    rpr.get_or_add_rFonts().set(qn('w:eastAsia'), '微软雅黑')

def p(text, bold=False, size=10.5, color=None, align=None, space_after=4):
    para = DOC.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    set_east_asia(run)
    para.paragraph_format.space_after = Pt(space_after)
    if align == 'center':
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return para

def h(level, text):
    para = DOC.add_heading(text, level=level)
    for run in para.runs:
        set_east_asia(run)
    return para

def bullets(items):
    for it in items:
        para = DOC.add_paragraph(style='List Bullet')
        run = para.add_run(it)
        run.font.size = Pt(10.5)
        set_east_asia(run)
        para.paragraph_format.space_after = Pt(2)

def mono(text, indent=0):
    para = DOC.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    rpr = run._element.get_or_add_rPr()
    rpr.get_or_add_rFonts().set(qn('w:eastAsia'), 'Consolas')
    para.paragraph_format.left_indent = Cm(0.6 + indent)
    para.paragraph_format.space_after = Pt(1)
    return para

def table(headers, rows, widths=None):
    t = DOC.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htxt in enumerate(headers):
        hdr[i].text = ''
        run = hdr[i].paragraphs[0].add_run(htxt)
        run.bold = True
        run.font.size = Pt(10)
        set_east_asia(run)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
            set_east_asia(run)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    DOC.add_paragraph()
    return t

# ================= 封面 =================
for _ in range(5):
    DOC.add_paragraph()
p('肺保护驾驶舱', bold=True, size=26, color='1F4E79', align='center', space_after=8)
p('—— 气道驱动压（ΔP）与气道机械功率（MP）累积暴露风险量化 + 趋势分析 ——', size=13, color='44546A', align='center', space_after=30)
p('系统设计说明书（SDD）', bold=True, size=16, align='center', space_after=6)
p('（第一阶段 · 普通机型 · 15寸触摸屏 WebUI）', size=12, align='center', space_after=40)
table(['项目', '内容'], [
    ['文档编号', 'SDD-VENT-DPMP-001'],
    ['版本', 'V1.0'],
    ['日期', '2026年8月26日'],
    ['技术栈', '前端 WebUI（HTML5/CSS3/JS）｜后端 Python 3.x（FastAPI，规划）｜数据库 MongoDB'],
    ['对应需求', 'SRS-VENT-DPMP-001 V1.0'],
    ['文档状态', '评审稿'],
], widths=[4, 11])
DOC.add_page_break()

# ================= 1 概述 =================
h(1, '1 设计概述')
h(2, '1.1 设计目标')
bullets([
    '在普通机型上以纯软件方式实现ΔP与MP的实时计算、累积暴露量化、趋势分析与风险可视化，不依赖任何额外硬件。',
    'ΔP与MP两模块合并为统一"肺保护驾驶舱"模块，共享数据管道、累积暴露分析引擎与风险可视化框架，一次搭建、双指标受益。',
    '界面简洁美观、触控友好，适配15寸触摸屏（1366×768），服务医生查房决策与护士日常监护两类核心工作流。',
    '架构上为第二、三阶段功能（跨肺压版本、P-SILI、WOB、ABG套件、AI指数统计预警）预留扩展接口，避免重复开发。',
])
h(2, '1.2 设计原则')
table(['原则', '说明'], [
    ['数据管道复用', 'ΔP/MP共享同一采集→计算→汇总→持久化→展示管道；后续功能仅替换"计算源"即可复用整套框架'],
    ['算法可插拔', '计算引擎以策略模式组织（公式、阈值、风险分级规则均配置化），新模式公式/新阈值不改动框架'],
    ['监测而非诊断', '全部输出定位为风险提示与决策支持，文案、交互、留痕设计均遵循该合规定位'],
    ['数据诚实', '不插值、不估算：断档标记为无数据；不可靠数据（强自主呼吸、阻断失败）显式标注并剔除统计'],
    ['触摸优先', '大热区、大字号、高对比度；核心路径≤2次触摸；3米可辨识风险色'],
], widths=[3.5, 11.5])
h(2, '1.3 系统边界')
bullets([
    '输入边界：呼吸机主控系统的气道压力、流量、容积信号与通气参数（经呼吸机MongoDB读取）。',
    '输出边界：设备本机15寸触摸屏WebUI（监测、趋势、风险图、预警、配置、导出）；数据写回MongoDB供导出与后续模块复用。',
    '本期范围：前端界面与数据契约完整交付；后端计算服务按本文档第5–8章设计分期实现，原型期以模拟数据驱动界面。',
])

# ================= 2 总体架构 =================
h(1, '2 总体架构')
h(2, '2.1 架构分层')
p('系统采用四层架构，部署于呼吸机整机嵌入式平台：')
table(['层', '组件', '职责', '技术选型'], [
    ['呈现层', 'WebUI 单页应用', '总览/趋势/二维风险图/预警中心/设置页面渲染与触控交互', 'HTML5 + CSS3 + 原生JS + SVG图表（设备内离线运行，不依赖外网）'],
    ['服务层', '数据服务 API + WebSocket推送', '指标查询、趋势查询、预警订阅、阈值配置、导出', 'Python 3.11 + FastAPI + Uvicorn（规划）'],
    ['计算层', '实时计算引擎 + 累积暴露引擎 + 趋势检测引擎', '逐呼ΔP/MP计算、1min汇总、AUC/占比/加权指数、CUSUM变化点、风险分级与预警判定', 'Python（NumPy/pandas），算法与Web服务同进程或独立进程（IPC）'],
    ['数据层', 'MongoDB', '原始波形参数读取 + 指标/事件/配置持久化', 'MongoDB 6.x，本机部署，WiredTiger'],
], widths=[1.8, 3.4, 5.8, 4.0])
p('数据流总览：', bold=True)
mono('呼吸机主控 → MongoDB(raw_data) → 实时计算引擎(逐呼ΔP/MP) → 累积暴露引擎(1min汇总/AUC/CUSUM)')
mono('→ MongoDB(metrics/events) → FastAPI(REST+WebSocket) → WebUI（总览/趋势/风险图/预警）')
mono('                                                            ↘ 预警引擎 → 预警中心 + 界面预警条')
h(2, '2.2 部署视图')
bullets([
    '呼吸机整机：嵌入式工控板（x86/ARM），运行Linux + 设备内浏览器内核（全屏WebUI）+ Python服务 + MongoDB，全部本机闭环，无外网依赖（满足医疗设备网络安全要求）。',
    'WebUI与Python服务通过本机HTTP/WebSocket（127.0.0.1）通信；WebUI静态资源打包于设备固件，随固件升级。',
    '预留：后续可通过医院内网（隔离VLAN）提供只读数据导出/科研数据服务（本期不实现，仅预留网络配置）。',
])

# ================= 3 数据管道 =================
h(1, '3 数据管道设计')
h(2, '3.1 管道阶段')
table(['阶段', '频率', '处理内容', '输出'], [
    ['S1 采集订阅', '逐呼吸周期（约0.2–4s）', '从MongoDB增量读取气道压/流量/容积波形与参数帧（Pplat、PEEP、PIP、VT、RR、模式、自主呼吸标志）', '原始周期帧'],
    ['S2 逐呼计算', '逐呼吸周期', 'ΔP=Pplat−PEEP_total；按模式选择MP公式；有效性判定（阻断成功、自主呼吸强度）', '逐呼ΔP/MP记录'],
    ['S3 分钟汇总', '每1min', '有效值均值/中位数/最大值、有效率；写入metrics_1min（持久化主表）', '1min指标文档'],
    ['S4 累积暴露', '每1min增量', '滚动更新各时间窗(1h/6h/24h/全程)的AUC、超阈值时间、占比、加权暴露指数、累积机械能', '累积指标（嵌入metrics_1min）'],
    ['S5 趋势检测', '每1min', 'CUSUM变化点检测；均值趋势斜率；生成趋势事件', '变化点事件'],
    ['S6 风险分级与预警', '每1min + 实时越限即时', '四级风险分级判定；预警触发/恢复；写入事件表并WebSocket推送', '预警事件'],
], widths=[2.6, 2.4, 7.0, 3.0])
h(2, '3.2 数据质量与状态机')
p('每个呼吸周期数据标注有效性状态机：')
mono('VALID(有效) → UNRELIABLE_DP(强自主呼吸，平台压不可靠，ΔP标注/剔除)')
mono('           → MP_SUSPENDED(强自主呼吸，MP暂停计算并提示)')
mono('           → INVALID(阻断失败/伪影/传感器异常，剔除全部统计)')
p('界面在任何聚合视图上同步显示"数据有效率"，断档时段在趋势图上以灰色区间标记"无数据"，禁止插值。')

# ================= 4 算法设计 =================
h(1, '4 算法设计（后端规划）')
h(2, '4.1 ΔP计算')
bullets([
    '吸气末阻断（inspiratory hold）获取Pplat；PEEP_total取呼气末正压（含内源性PEEP时取实测呼气末压）。',
    'ΔP = Pplat − PEEP_total（cmH₂O）。',
    '触发条件：控制/被动通气下按固定间隔（默认每4次呼吸执行一次阻断，可配置）自动执行；强自主呼吸状态下不执行并标注不可靠。',
    '逐呼插值：两次阻断之间以流量-容积法估算瞬时驱动压作为过渡显示值（仅显示用，不进入累积统计；统计仅用实测阻断值，保证循证口径一致）。',
])
h(2, '4.2 MP计算（模式自适应）')
table(['通气模式', '公式', '说明'], [
    ['VCV', 'MP = 0.098 × RR × VT × (PIP − 0.5×ΔP)', 'Gattinoni简化公式（J/min，VT单位L）'],
    ['PCV', 'MP = 0.098 × RR × VT × (ΔP + PEEP)', 'Becher公式'],
    ['快速筛查参考', 'MP ≈ 4×ΔP + RR', 'Costa线性简化（设置中可选显示，仅参考）'],
], widths=[3, 7, 5])
bullets([
    '金标准为几何法（P-V环吸气支面积×RR），本期以简化公式实现，几何法作为二阶段校准升级项。',
    '标准化参考：按PBW显示标准化值（MP/PBW, J/min/kg，P1展示项）；Crs标准化（RAMPS框架：MP×时间×顺应性）作为后续研究扩展。',
])
h(2, '4.3 累积暴露量化引擎')
p('以1分钟粒度序列 {t_i, x_i, valid_i}（x为ΔP或MP均值）与阈值T计算：')
bullets([
    '超阈值累积时间 TimeAbove = Σ max(0, x_i − T) > 0 ? 1min : 0（分段线性插值边界可精确到秒，本期分钟粒度即可）。',
    '超阈值曲线下面积 AUC_above = Σ max(0, x_i − T) × Δt（ΔP单位 cmH₂O·h；MP单位 J）。',
    '时间占比 RatioAbove = TimeAbove / 窗口有效监测时间。',
    '加权暴露指数 WEI = Σ max(0, x_i − T)² × Δt / 窗口时长（对越限幅度平方加权，突出重度越限）。',
    '累积机械能 CumEnergy = Σ MP_i × Δt（不设阈值，全程累积）。',
    '时间窗：1h / 6h / 24h / 全程，滚动计算，增量更新（O(1)滑动窗口）。',
])
h(2, '4.4 趋势检测：CUSUM变化点')
bullets([
    '对1min ΔP/MP序列计算双边CUSUM：C⁺_t = max(0, C⁺_{t−1} + (x_t − μ₀ − k))，C⁻对称；H=5σ/k越界即报变化点。',
    '变化点在趋势图上以标记呈现（"风险拐点提示"），本期不联动报警（P1再评估）。',
    '变化点事件附带前后均值对比（如 ΔP 12→17 cmH₂O），帮助医生定位参数调整/病情变化时刻。',
])
h(2, '4.5 风险分级规则（配置化）')
table(['等级', '颜色', '判定规则（示例，均可配置）'], [
    ['低风险', '绿', 'ΔP<15 且 MP15min均值<17 且 24h超阈值占比<5%'],
    ['关注', '黄', '任一指标接近阈值（ΔP≥13或MP≥15）或1h超阈值占比≥5%'],
    ['累积暴露偏高', '橙', '实时越限持续≥5min，或1h占比≥10%，或24h占比≥15%，或WEI达橙色线'],
    ['高风险', '红', 'ΔP≥20 或 MP≥25 持续≥5min，或24h占比≥30%'],
], widths=[2.8, 1.6, 10.6])
p('分级输出同时驱动：界面风险色带、预警中心、日报卡。分级规则以JSON配置存储于MongoDB（config集合），修改无需发版。')

# ================= 5 数据库设计 =================
h(1, '5 MongoDB数据模型设计')
h(2, '5.1 集合设计（本模块新增）')
p('（1）metrics_1min —— 1分钟指标主表（核心持久化表）')
mono('{')
mono('  "_id": ObjectId,')
mono('  "patient_id": "P20260826-001", "vent_session_id": "V20260826-01",', 1)
mono('  "ts": ISODate("2026-08-26T14:00:00Z"),          // 分钟起始时间', 1)
mono('  "dp":  {"mean": 14.2, "median": 14.0, "max": 16.1, "pplat": 24.3, "peep": 10.1,')
mono('          "valid_ratio": 1.0, "status": "VALID"},', 2)
mono('  "mp":  {"mean": 16.8, "j_per_min": true, "formula": "VCV_GATTINONI",')
mono('          "mode": "VCV", "rr": 22, "vt_ml": 420, "pip": 28.4, "valid_ratio": 1.0},', 2)
mono('  "cum": {                                            // 累积暴露（各时间窗滚动）', 1)
mono('          "dp_auc_1h": 0.8, "dp_time_above_1h": 12,  // 单位: cmH2O·h, min', 2)
mono('          "dp_ratio_1h": 0.20, "dp_wei_1h": 0.4,', 2)
mono('          "dp_auc_24h": 6.2, "dp_ratio_24h": 0.18, ...', 2)
mono('          "mp_energy_24h": 23400, "mp_ratio_24h": 0.12, ... },', 2)
mono('  "risk_level": "ORANGE", "chg_pt": {"dp": false, "mp": true}', 1)
mono('}')
p('索引：{patient_id, ts:-1} 唯一；{ts} TTL归档（默认保留90天，可配置）。')
p('（2）events —— 预警与事件表')
mono('{ "_id": ObjectId, "patient_id": "...", "type": "DP_THRESHOLD_EXCEEDED",')
mono('  "level": "ORANGE", "start_ts": ISODate, "end_ts": ISODate|null,')
mono('  "detail": {"dp": 16.2, "threshold": 15, "duration_min": 12},')
mono('  "ack": {"by": "nurse_li", "ts": ISODate, "action": "acknowledged"} }')
p('（3）threshold_config —— 阈值与规则配置（带审计字段）')
mono('{ "dp_threshold": 15, "mp_threshold": 17, "rules": {risk JSON},')
mono('  "updated_by": "dr_wang", "updated_at": ISODate, "prev": {...} }')
p('（4）read_raw 视图 —— 对呼吸机原始库的只读映射（不重复存储，服务层只读账号访问）。')
h(2, '5.2 与既有数据的关系')
bullets([
    '原始波形与参数由呼吸机主控写入其既有库，本模块以只读账号访问，职责边界清晰。',
    '本模块产出的metrics_1min/events同时是后续第二阶段功能（P-SILI、WOB）与AI指数统计预警的输入源，一次沉淀、多模块复用。',
    '导出功能直接对metrics_1min做CSV投影，保证质控/科研口径与设备界面一致。',
])

# ================= 6 API设计 =================
h(1, '6 后端API设计（REST + WebSocket）')
table(['方法/协议', '路径', '说明'], [
    ['GET', '/api/v1/overview?patient_id=', '总览页数据包：当前值、1m/15m均值、各窗累积指标、风险等级、活动预警'],
    ['GET', '/api/v1/trends?metric=dp|mp&window=1h|6h|24h&from=&to=', '趋势序列（1min粒度）+ 阈值线 + 超阈值区间 + 变化点标记 + 断档区间'],
    ['GET', '/api/v1/riskmap?window=1h', '二维风险图数据：最近1h轨迹点列（ts、dp、mp、风险等级）'],
    ['GET', '/api/v1/events?level=&from=&to=&status=', '预警/事件列表（分页）'],
    ['POST', '/api/v1/events/{id}/ack', '预警确认/忽略（留痕）'],
    ['GET / PUT', '/api/v1/config/thresholds', '读取/修改阈值与规则（PUT需权限+二次确认语义，自动写审计）'],
    ['GET', '/api/v1/export/csv?metric=&from=&to=', 'CSV导出（metrics_1min投影）'],
    ['WebSocket', '/ws/realtime', '实时推送：逐呼ΔP/MP（节流至1s）、风险等级变更、预警触发/恢复、计算状态（MP_SUSPENDED等）'],
], widths=[2.2, 6.8, 6.0])
p('约定：JSON字段名与5.1数据模型一致；时间统一ISO8601+UTC；错误码遵循设备内服务统一规范；接口幂等、只读接口无副作用。')

# ================= 7 前端设计 =================
h(1, '7 前端（WebUI）设计')
h(2, '7.1 页面结构与导航')
table(['页面', '内容', '对应需求'], [
    ['① 总览（默认页）', '整体风险评级横幅（等级大字+一句话结论）+ ΔP/MP仪表盘卡（半圆仪表盘：绿/黄/橙/红分区弧+阈值刻度+当前值标记+中央大数字，配一句话提示，一图一数一句话原则）+ 能量-应力二维风险图卡（近1h轨迹+象限结论）+ 24h风险分级时间带 + 通气参数快照', 'FR-A1/A2/B1/B2/C2/D1'],
    ['② ΔP暴露趋势', '24h(可切1h/6h)ΔP均值趋势 + 15阈值线 + 超阈值区域红色填充高亮 + 变化点标记 + 断档灰区 + 累积指标侧栏', 'FR-B1/C1'],
    ['③ MP暴露趋势', '同②，MP版本 + 累积机械能曲线（面积图）+ 公式/模式标注', 'FR-B2'],
    ['④ 能量-应力风险图', '横轴ΔP/纵轴MP二维图，四象限+阈值参考线+1h轨迹拖尾，点按查看时刻详情', 'FR-D1'],
    ['⑤ 预警中心', '事件列表（等级/时间/时长/状态），确认/忽略，历史筛选', 'FR-C2'],
    ['⑥ 设置', '阈值调整（二次确认+审计）、公式选项、显示选项、帮助/局限性说明', 'FR-E1/E3'],
], widths=[2.8, 9.6, 2.6])
p('导航：左侧固定竖排大图标页签（触控热区≥56px），总览页任何核心信息≤1次触摸下钻到对应详情页。')
h(2, '7.2 15寸触摸屏交互规范')
table(['项', '规范'], [
    ['分辨率/缩放', '设计基准1366×768；使用viewport缩放与rem布局，兼容1280×800等备用屏'],
    ['触控热区', '普通按钮≥44×44px，核心按钮（页签/确认）≥56×56px；控件间距≥8px防误触'],
    ['字号', '主数值≥40px（3米可辨），标签≥16px，正文≥14px，最小注释12px'],
    ['色彩语义', '绿#2E9E5B 低风险｜黄#E8B931 关注｜橙#E67E22 累积偏高｜红#E74C3C 高风险；背景浅色为主（医疗环境、昼夜可读），风险色仅用于数据与告警，不做装饰'],
    ['反馈', '触摸即时高亮反馈；数据刷新有淡入；预警触发顶部横幅+色带变化+（可选）声音，声音可由护士静音并留痕'],
    ['防误操作', '阈值修改：滑条+数值确认+二次弹窗；预警忽略：长按0.5s确认'],
], widths=[3, 12])
h(2, '7.3 图表实现')
bullets([
    '设备内离线运行：图表以原生SVG/Canvas自绘（趋势折线、面积填充、二维散点轨迹），不引入外部CDN依赖，保证固件离线可用与安全审计可控。',
    '趋势图交互：双指/捏合或按钮缩放时间窗，长按游标查看逐分钟数值；超阈值区域自动红色半透明填充并生成"超阈值时段"列表。',
    '二维风险图：轨迹以时间渐变（新亮旧暗），当前点脉冲动画；象限文案常驻小字，避免遮挡数据。',
])
h(2, '7.4 前端状态与实时性')
bullets([
    '单页应用，路由五页；状态机管理 WebSocket 实时数据 → 视图模型 → 组件渲染，断线自动重连并以最后数据时间戳提示"数据延迟"。',
    '总览页实时值1s节流刷新；趋势页按窗口节流（30s）；页面切换预取下一页首屏数据，保证≤1s切换。',
])

# ================= 8 可靠性安全合规 =================
h(1, '8 可靠性、安全与合规设计')
bullets([
    '进程守护：Python计算服务与Web服务由设备init系统守护，崩溃自动重启；重启后从metrics_1min恢复累积状态（幂等重算最近窗口）。',
    '数据完整性：分钟汇总写入采用upsert（patient_id+ts唯一键），断电恢复后补算缺口并标注。',
    '权限：本机操作分级（护士/医师/工程师），阈值修改需医师级权限；所有配置变更写审计。',
    '合规定位：界面文案统一为"风险提示/监测/决策支持"，禁用诊断性表述；帮助页明示循证依据与局限性；为后续SaMD II类注册预留文档与数据追溯链（需求-设计-代码-用例对应）。',
    '网络安全：服务仅绑定本机回环；无外网通信；导出功能需物理USB介质+权限确认。',
])

# ================= 9 测试策略 =================
h(1, '9 测试策略（概要）')
table(['类别', '要点'], [
    ['算法单元测试', '标准波形回放（含VCV/PCV/模式切换/强自主呼吸/阻断失败用例）验证ΔP/MP误差≤±2%；AUC/占比/WEI/累积机械能与手工基准一致'],
    ['引擎测试', '1min汇总幂等；滑动窗口滚动正确；断电恢复补算；TTL归档'],
    ['前端测试', '1366×768全页面走查；触控热区/字号规范核查；风险色一致性；断线/数据延迟/断档显示'],
    ['预警测试', '四级分级规则全组合触发/恢复用例；确认留痕'],
    ['性能测试', '24h趋势查询≤2s（约1440点）；页面切换≤1s；WebSocket推送延迟≤1s'],
    ['回归与追溯', '用例编号映射SRS需求编号（FR-*），实现需求-用例双向追溯矩阵'],
], widths=[3.5, 11.5])

# ================= 10 迭代规划 =================
h(1, '10 实施与迭代规划')
h(2, '10.1 本期（原型+一期开发）')
table(['迭代', '内容', '产出'], [
    ['M1 原型', 'WebUI全页面高保真原型（模拟数据），数据契约冻结', '本文档+原型（已交付）'],
    ['M2 数据管道', 'S1采集订阅+S2逐呼计算+metrics_1min落库（离线回放驱动）', '可回放验证的计算引擎'],
    ['M3 累积与趋势', 'S4/S5累积暴露引擎+CUSUM；趋势页接真实数据', '趋势功能可用'],
    ['M4 风险与预警', 'S6分级规则+预警中心+WebSocket实时推送', '模块端到端可用'],
    ['M5 联调验收', '整机联调、性能、可靠性、验收测试', '一期发布'],
], widths=[2.6, 8.4, 4.0])
h(2, '10.2 后续阶段衔接（共享框架扩展点）')
table(['后续功能', '复用点', '扩展点'], [
    ['跨肺压版本（高端机型）', '整套管道/累积引擎/趋势与风险可视化框架', '替换驱动压计算源为食管压版本；增加信号质量指数；双视图切换'],
    ['P-SILI / WOB（第二阶段）', '食管压信号平台、趋势展示框架、事件/预警机制', '新增信号处理与六分量算法，接入统一风险概览'],
    ['AI指数统计预警', 'metrics_1min多模块指标沉淀', '归一化统一评分引擎，作为"总控"输出'],
    ['ABG套件（第三阶段）', '数据契约、导出、日报卡', '血气录入/无创趋势接入，力学+血气联合分析'],
], widths=[3.6, 6.4, 5.0])
p('说明：本期原型已按上述数据契约以模拟数据实现界面交互，后端各服务按M2–M4迭代落地。', size=9.5, color='808080')

DOC.save('设计文档-肺保护驾驶舱（ΔP与MP累积暴露）.docx')
print('SAVED: 设计文档')
